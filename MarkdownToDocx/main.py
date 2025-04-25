import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

def add_html_to_docx(soup, doc):
    for element in soup.contents:
        if element.name and element.name.startswith('h') and len(element.name) == 2 and element.name[1].isdigit():
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)

        elif element.name == 'p':
            para = doc.add_paragraph()
            for child in element.children:
                if isinstance(child, str):
                    para.add_run(child)
                elif child.name in ['strong', 'b']:
                    run = para.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['em', 'i']:
                    run = para.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = para.add_run(child.get_text())
                    run.font.name = 'Courier New'
                else:
                    para.add_run(child.get_text())

        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='ListBullet')

        elif element.name == 'ol':
            doc.add_paragraph()
            count = 1
            for li in element.find_all('li', recursive=False):
                para = doc.add_paragraph(f"{count}. {li.get_text()}")
                count += 1
            doc.add_paragraph()


        elif element.name == 'blockquote':
            para = doc.add_paragraph(element.get_text())
            para.paragraph_format.left_indent = Pt(20)

def markdown_to_docx(md_file, docx_file):
    # Read Markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Convert Markdown to HTML
    html = markdown.markdown(md_text, extensions=['extra'])

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')

    # Create DOCX
    doc = Document()
    add_html_to_docx(soup, doc)

    # Save the DOCX file
    doc.save(docx_file)
    print(f"Converted '{md_file}' to '{docx_file}' successfully.")

# Example usage
markdown_to_docx('markdown.md', 'output.docx')
